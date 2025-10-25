from docx import Document
from datetime import datetime
import os

TAX_TYPE_CGST_SGST = "CGST_SGST"
TAX_TYPE_IGST = "IGST"


def calculate_totals(items, discount=0.0, tax_type=TAX_TYPE_CGST_SGST, gst_rate=18):
    subtotal = sum(i["qty"] * i["rate"] for i in items)

    # Apply discount
    discount_amt = (discount / 100) * subtotal if discount > 0 else 0
    taxable_value = subtotal - discount_amt

    # Tax split logic
    if tax_type == TAX_TYPE_CGST_SGST:
        cgst = sgst = taxable_value * (gst_rate / 2) / 100
        igst = 0
    else:
        cgst = sgst = 0
        igst = taxable_value * gst_rate / 100

    total = taxable_value + cgst + sgst + igst

    return {
        "subtotal": round(subtotal, 2),
        "discount_amt": round(discount_amt, 2),
        "taxable_value": round(taxable_value, 2),
        "cgst": round(cgst, 2),
        "sgst": round(sgst, 2),
        "igst": round(igst, 2),
        "total": round(total, 2),
    }


def generate_invoice(template_path, output_path, invoice_data):
    doc = Document(template_path)

    # Replace text placeholders
    merge_fields = {
        "[INVOICE_NO]": invoice_data["invoice_no"],
        "[DATE]": invoice_data.get("date", datetime.now().strftime("%d-%m-%Y")),
        "[GST_NUMBER]": invoice_data.get("gst_number", ""),
        "[STATE]": invoice_data.get("state", ""),
        "[YES/NO]": invoice_data.get("reverse_charge", "NO"),
        "[TOTAL_QTY]": str(sum(i["qty"] for i in invoice_data["items"])),
        "[TOTAL_AMOUNT]": str(invoice_data["totals"]["total"]),
        "[AMOUNT_IN_WORDS]": invoice_data.get("amount_in_words", ""),
    }

    for p in doc.paragraphs:
        for key, value in merge_fields.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in merge_fields.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Fill table
    table = doc.tables[2]  # ITEM TABLE is table 2
    gst_rate = invoice_data.get("gst_rate", 18)
    for row_data in invoice_data["items"]:
        row_cells = table.add_row().cells
        total = row_data["qty"] * row_data["rate"]
        taxable = total  # Assuming no discount per item for simplicity
        cgst_amt = taxable * (gst_rate / 2) / 100 if invoice_data["tax_type"] == TAX_TYPE_CGST_SGST else 0
        sgst_amt = cgst_amt
        igst_amt = taxable * gst_rate / 100 if invoice_data["tax_type"] == TAX_TYPE_IGST else 0
        net_amt = taxable + cgst_amt + sgst_amt + igst_amt
        row_cells[0].text = str(row_data["sl"])
        row_cells[1].text = row_data["description"]
        row_cells[2].text = str(row_data["hsn"])
        row_cells[3].text = str(row_data["qty"])
        row_cells[4].text = ""  # Unit, assuming empty or not provided
        row_cells[5].text = f"{row_data['rate']:.2f}"
        row_cells[6].text = f"{total:.2f}"
        row_cells[7].text = ""  # Disc %, assuming 0 or not provided
        row_cells[8].text = ""  # Disc Amt, assuming 0 or not provided
        row_cells[9].text = f"{taxable:.2f}"
        row_cells[10].text = f"{gst_rate/2:.1f}" if invoice_data["tax_type"] == TAX_TYPE_CGST_SGST else ""
        row_cells[11].text = f"{cgst_amt:.2f}" if cgst_amt else ""
        row_cells[12].text = f"{gst_rate/2:.1f}" if invoice_data["tax_type"] == TAX_TYPE_CGST_SGST else ""
        row_cells[13].text = f"{sgst_amt:.2f}" if sgst_amt else ""
        row_cells[14].text = f"{net_amt:.2f}"

    totals = invoice_data["totals"]
    gst_rate = invoice_data.get("gst_rate", 18)

    # Replace totals table placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in {
                    "<<SUBTOTAL>>": totals["subtotal"],
                    "<<DISCOUNT>>": totals["discount_amt"],
                    "<<TAXABLE_VALUE>>": totals["taxable_value"],
                    "<<CGST>>": totals["cgst"],
                    "<<SGST>>": totals["sgst"],
                    "<<IGST>>": totals["igst"],
                    "<<TOTAL>>": totals["total"],
                }.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Replace tax summary table placeholders
    for table in doc.tables:
        for row in table.rows:
            if any("[ ]" in cell.text for cell in row.cells):
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = cells[0].text.replace("[ ]", str(gst_rate))
                    cells[1].text = cells[1].text.replace("[ ]", str(totals["taxable_value"]))
                    cells[2].text = cells[2].text.replace("[ ]", str(totals["cgst"]))
                    cells[3].text = cells[3].text.replace("[ ]", str(totals["sgst"]))
                    cells[4].text = cells[4].text.replace("[ ]", str(round(totals["cgst"] + totals["sgst"] + totals["igst"], 2)))
                break

    doc.save(output_path)
    return output_path


# ✅ Example Usage
if __name__ == "__main__":
    items = [
        {"sl": 1, "description": "Product A", "hsn": "1001", "qty": 2, "rate": 150},
        {"sl": 2, "description": "Product B", "hsn": "2002", "qty": 1, "rate": 500},
    ]

    totals = calculate_totals(items, discount=5, tax_type=TAX_TYPE_CGST_SGST)

    invoice_data = {
        "invoice_no": "INV-001",
        "client_name": "Client Name Here",
        "client_address": "Client Address, City",
        "items": items,
        "discount": 5,
        "tax_type": TAX_TYPE_CGST_SGST,
        "totals": totals,
        "amount_in_words": "Eight Hundred Ninety Six and Eighty Paise Only",
    }

    output = generate_invoice(
        "Invoice_Template.docx",
        "Generated_Invoice.docx",
        invoice_data
    )

    print(f"✅ Invoice generated: {output}")
